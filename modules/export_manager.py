"""Export functionality for research outputs"""

import os
from datetime import datetime
from typing import Dict, Any, List
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import json

class ExportManager:
    """Manage export of research results to various formats"""
    
    def __init__(self, export_path: str = "./exports"):
        self.export_path = export_path
        self._ensure_export_directory()
    
    def _ensure_export_directory(self):
        """Ensure export directory exists"""
        if not os.path.exists(self.export_path):
            os.makedirs(self.export_path)
    
    def export_to_pdf(
        self,
        content: Dict[str, Any],
        filename: str = None
    ) -> str:
        """
        Export content to PDF format
        
        Args:
            content: Content to export
            filename: Output filename
            
        Returns:
            Path to exported file
        """
        if not filename:
            filename = f"research_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        filepath = os.path.join(self.export_path, filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1E3A8A'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        title = content.get('title', 'Research Export')
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # Add metadata
        if 'metadata' in content:
            meta_style = ParagraphStyle(
                'MetaData',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#64748B')
            )
            
            for key, value in content['metadata'].items():
                story.append(Paragraph(f"<b>{key}:</b> {value}", meta_style))
            story.append(Spacer(1, 20))
        
        # Add main content
        if 'full_review' in content:
            content_style = ParagraphStyle(
                'ContentStyle',
                parent=styles['Normal'],
                fontSize=11,
                leading=14
            )
            
            # Split content into paragraphs
            paragraphs = content['full_review'].split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.replace('\n', '<br/>'), content_style))
                    story.append(Spacer(1, 12))
        
        # Add sections if available
        if 'sections' in content:
            for section_name, section_content in content['sections'].items():
                if section_content:
                    # Section heading
                    section_style = ParagraphStyle(
                        'SectionHeading',
                        parent=styles['Heading2'],
                        fontSize=16,
                        textColor=colors.HexColor('#3B82F6'),
                        spaceBefore=20,
                        spaceAfter=10
                    )
                    
                    story.append(Paragraph(section_name.replace('_', ' ').title(), section_style))
                    story.append(Paragraph(section_content.replace('\n', '<br/>'), content_style))
                    story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        
        return filepath
    
    def export_to_word(
        self,
        content: Dict[str, Any],
        filename: str = None
    ) -> str:
        """
        Export content to Word document
        
        Args:
            content: Content to export
            filename: Output filename
            
        Returns:
            Path to exported file
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if not filename:
            filename = f"research_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        filepath = os.path.join(self.export_path, filename)
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(content.get('title', 'Research Export'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        if 'metadata' in content:
            doc.add_heading('Document Information', level=1)
            for key, value in content['metadata'].items():
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))
        
        # Add main content
        if 'full_review' in content:
            doc.add_heading('Literature Review', level=1)
            paragraphs = content['full_review'].split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
        
        # Add sections
        if 'sections' in content:
            for section_name, section_content in content['sections'].items():
                if section_content:
                    doc.add_heading(section_name.replace('_', ' ').title(), level=2)
                    doc.add_paragraph(section_content)
        
        # Save document
        doc.save(filepath)
        
        return filepath
    
    def export_to_markdown(
        self,
        content: Dict[str, Any],
        filename: str = None
    ) -> str:
        """
        Export content to Markdown format
        
        Args:
            content: Content to export
            filename: Output filename
            
        Returns:
            Path to exported file
        """
        if not filename:
            filename = f"research_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        
        filepath = os.path.join(self.export_path, filename)
        
        markdown_content = []
        
        # Add title
        markdown_content.append(f"# {content.get('title', 'Research Export')}\n")
        
        # Add metadata
        if 'metadata' in content:
            markdown_content.append("## Document Information\n")
            for key, value in content['metadata'].items():
                markdown_content.append(f"- **{key}**: {value}")
            markdown_content.append("")
        
        # Add main content
        if 'full_review' in content:
            markdown_content.append("## Literature Review\n")
            markdown_content.append(content['full_review'])
            markdown_content.append("")
        
        # Add sections
        if 'sections' in content:
            for section_name, section_content in content['sections'].items():
                if section_content:
                    markdown_content.append(f"## {section_name.replace('_', ' ').title()}\n")
                    markdown_content.append(section_content)
                    markdown_content.append("")
        
        # Write to file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_content))
        
        return filepath
    
    def export_to_latex(
        self,
        content: Dict[str, Any],
        filename: str = None
    ) -> str:
        """
        Export content to LaTeX format
        
        Args:
            content: Content to export
            filename: Output filename
            
        Returns:
            Path to exported file
        """
        if not filename:
            filename = f"research_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.tex"
        
        filepath = os.path.join(self.export_path, filename)
        
        latex_content = []
        
        # LaTeX preamble
        latex_content.extend([
            "\\documentclass[12pt]{article}",
            "\\usepackage[utf8]{inputenc}",
            "\\usepackage{geometry}",
            "\\geometry{a4paper, margin=1in}",
            "\\usepackage{hyperref}",
            "\\usepackage{graphicx}",
            "",
            f"\\title{{{content.get('title', 'Research Export')}}}",
            f"\\author{{IntelliDoc Research Pro}}",
            f"\\date{{\\today}}",
            "",
            "\\begin{document}",
            "\\maketitle",
            ""
        ])
        
        # Add abstract if available
        if 'executive_summary' in content.get('sections', {}):
            latex_content.extend([
                "\\begin{abstract}",
                self._escape_latex(content['sections']['executive_summary']),
                "\\end{abstract}",
                ""
            ])
        
        # Add main content
        if 'full_review' in content:
            latex_content.append("\\section{Literature Review}")
            latex_content.append(self._escape_latex(content['full_review']))
            latex_content.append("")
        
        # Add sections
        if 'sections' in content:
            for section_name, section_content in content['sections'].items():
                if section_content and section_name != 'executive_summary':
                    section_title = section_name.replace('_', ' ').title()
                    latex_content.append(f"\\section{{{section_title}}}")
                    latex_content.append(self._escape_latex(section_content))
                    latex_content.append("")
        
        latex_content.append("\\end{document}")
        
        # Write to file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(latex_content))
        
        return filepath
    
    def export_to_json(
        self,
        content: Dict[str, Any],
        filename: str = None
    ) -> str:
        """
        Export content to JSON format
        
        Args:
            content: Content to export
            filename: Output filename
            
        Returns:
            Path to exported file
        """
        if not filename:
            filename = f"research_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        filepath = os.path.join(self.export_path, filename)
        
        # Add export timestamp
        export_data = {
            "export_timestamp": datetime.now().isoformat(),
            "export_version": "1.0.0",
            **content
        }
        
        # Write to file
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        return filepath
    
    def export_bibliography(
        self,
        citations: List[str],
        format: str = "txt",
        filename: str = None
    ) -> str:
        """
        Export bibliography in various formats
        
        Args:
            citations: List of formatted citations
            format: Export format (txt, bib, ris)
            filename: Output filename
            
        Returns:
            Path to exported file
        """
        if not filename:
            filename = f"bibliography_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{format}"
        
        filepath = os.path.join(self.export_path, filename)
        
        if format == "txt":
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("Bibliography\n")
                f.write("=" * 50 + "\n\n")
                for i, citation in enumerate(citations, 1):
                    f.write(f"[{i}] {citation}\n\n")
        
        elif format == "bib":
            # BibTeX format
            with open(filepath, 'w', encoding='utf-8') as f:
                for i, citation in enumerate(citations, 1):
                    f.write(f"@article{{ref{i},\n")
                    f.write(f"  title = {{{citation}}},\n")
                    f.write(f"  year = {{2024}}\n")
                    f.write("}\n\n")
        
        return filepath
    
    def _escape_latex(self, text: str) -> str:
        """
        Escape special LaTeX characters
        
        Args:
            text: Text to escape
            
        Returns:
            Escaped text
        """
        replacements = {
            '\\': '\\textbackslash{}',
            '{': '\\{',
            '}': '\\}',
            '$': '\\$',
            '&': '\\&',
            '%': '\\%',
            '#': '\\#',
            '_': '\\_',
            '~': '\\textasciitilde{}',
            '^': '\\textasciicircum{}'
        }
        
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        
        return text